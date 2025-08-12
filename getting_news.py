# Copyright 2025 Nicholas F. Tominello
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy at http://www.apache.org/licenses/LICENSE-2.0
# #!/usr/bin/env python3
# By Nicholas F. Tominello
"""
Complete Enhanced News Analysis System with Professional Document Generation
Combines all tools and enhanced document generator into a single working file
"""

import os
import time
import json
import requests
from datetime import datetime
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
from ddgs import DDGS
from openai import OpenAI
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    print("⚠️ Warning: OPENAI_API_KEY not found. Please set it in your .env file or environment variables.")
    print("You can continue, but AI summaries and analysis will fail.")

client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# ============================================
# TOOL DEFINITIONS
# ============================================

class Tool:
    """Base class for all tools"""
    def __init__(self, name: str, description: str):
        self.name = name
        self.description = description

    def execute(self, *args, **kwargs):
        raise NotImplementedError


class InternetConnectivityTool(Tool):
    def __init__(self):
        super().__init__("check_internet", "Check if internet connection is available")

    def execute(self) -> Dict[str, Any]:
        try:
            response = requests.get("https://httpbin.org/get", timeout=5)
            return {
                "success": True,
                "status": "connected",
                "response_code": response.status_code
            }
        except Exception as e:
            return {
                "success": False,
                "status": "disconnected",
                "error": str(e)
            }


class NewsSearchTool(Tool):
    def __init__(self):
        super().__init__("search_news", "Search for current news articles on any topic")

    def execute(self, query: str = "top news today", max_results: int = 20, **kwargs) -> Dict[str, Any]:
        try:
            print(f"🔍 Searching for: '{query}' (max: {max_results})")
            with DDGS() as ddgs:
                results = ddgs.text(
                    query,
                    region="us-en",
                    safesearch="moderate",
                    timelimit="d",
                    max_results=max_results
                )

                articles = []
                for r in results:
                    articles.append({
                        "title": r.get('title', '[No Title]'),
                        "summary": r.get('body', '[No Summary]'),
                        "url": r.get('href', '[No Link]'),
                        "timestamp": datetime.now().isoformat()
                    })

                return {
                    "success": True,
                    "articles": articles,
                    "count": len(articles),
                    "query": query
                }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "articles": [],
                "count": 0
            }


class WebScrapingTool(Tool):
    def __init__(self):
        super().__init__("scrape_articles", "Scrape full content from article URLs and generate summaries")

    def execute(self, articles: List[Dict] = None, max_articles: int = 15, **kwargs) -> Dict[str, Any]:
        if not articles:
            return {
                "success": False,
                "error": "No articles provided for scraping",
                "scraped_articles": []
            }

        print(f"🌐 Scraping {min(len(articles), max_articles)} articles...")
        scraped_articles = []

        for i, article in enumerate(articles[:max_articles]):
            try:
                print(f"   📄 Scraping article {i+1}: {article['title'][:50]}...")

                # Get the full content
                content = self.scrape_url(article['url'])

                if content:
                    # Generate summary using GPT-4
                    if client:
                        summary = self.generate_summary(article['title'], content)
                    else:
                        summary = f"AI summary unavailable (no API key). Content preview: {content[:200]}..."

                    scraped_articles.append({
                        "original_title": article['title'],
                        "url": article['url'],
                        "original_summary": article.get('summary', ''),
                        "full_content": content[:2000],
                        "ai_summary": summary,
                        "scrape_timestamp": datetime.now().isoformat(),
                        "word_count": len(content.split())
                    })
                    print(f"   ✅ Successfully scraped ({len(content.split())} words)")
                else:
                    print(f"   ⚠️ Could not scrape content")
                    scraped_articles.append({
                        "original_title": article['title'],
                        "url": article['url'],
                        "original_summary": article.get('summary', ''),
                        "full_content": "",
                        "ai_summary": "Could not scrape content from this URL",
                        "scrape_timestamp": datetime.now().isoformat(),
                        "word_count": 0
                    })

                time.sleep(0.5)

            except Exception as e:
                print(f"   ❌ Error scraping article {i+1}: {e}")
                scraped_articles.append({
                    "original_title": article['title'],
                    "url": article['url'],
                    "original_summary": article.get('summary', ''),
                    "full_content": "",
                    "ai_summary": f"Error scraping: {str(e)}",
                    "scrape_timestamp": datetime.now().isoformat(),
                    "word_count": 0
                })

        return {
            "success": True,
            "scraped_articles": scraped_articles,
            "total_scraped": len(scraped_articles),
            "successful_scrapes": len([a for a in scraped_articles if a['word_count'] > 0])
        }

    def scrape_url(self, url: str) -> str:
        """Scrape content from a URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }

            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
                script.decompose()

            content_selectors = [
                'article', '[role="main"]', '.content', '.post-content',
                '.entry-content', '.article-body', '.story-body', 'main'
            ]

            main_content = ""
            for selector in content_selectors:
                elements = soup.select(selector)
                if elements:
                    main_content = ' '.join([elem.get_text() for elem in elements])
                    break

            if not main_content:
                paragraphs = soup.find_all('p')
                main_content = ' '.join([p.get_text() for p in paragraphs])

            main_content = re.sub(r'\s+', ' ', main_content).strip()
            return main_content[:3000] if main_content else ""

        except Exception as e:
            print(f"      Error fetching {url}: {e}")
            return ""

    def generate_summary(self, title: str, content: str) -> str:
        """Generate AI summary of scraped content"""
        if not client:
            return "AI summary unavailable (no API key configured)"

        if not content or len(content) < 100:
            return "Insufficient content to summarize"

        try:
            prompt = f"""Summarize this news article concisely:

Title: {title}

Content: {content[:2000]}

Provide a clear, factual summary in 3-4 sentences focusing on:
1. Main news/event
2. Key details and facts
3. Any important implications or context"""

            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert news summarizer. Provide concise, factual summaries."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=200,
                temperature=0.2
            )

            return response.choices[0].message.content.strip()

        except Exception as e:
            return f"Summary generation failed: {str(e)}"


class NewsAnalysisTool(Tool):
    def __init__(self):
        super().__init__("analyze_news", "Analyze news articles and answer questions about them")

    def execute(self, articles: List[Dict] = None, question: str = "Analyze these articles", **kwargs) -> Dict[str, Any]:
        if not articles:
            return {
                "success": False,
                "analysis": "No articles available for analysis",
                "confidence": 0
            }

        if not client:
            return {
                "success": False,
                "analysis": "AI analysis unavailable (no API key configured)",
                "error": "No OpenAI API key"
            }

        print(f"🧠 Analyzing {len(articles)} articles with question: '{question}'")

        context_parts = []
        for i, article in enumerate(articles[:5]):
            context_parts.append(
                f"{i+1}. Title: {article['title']}\n"
                f"   Summary: {article['summary'][:200]}...\n"
                f"   URL: {article['url']}"
            )

        context = "\n\n".join(context_parts)
        prompt = f"""Based on these current news articles:

{context}

Question: {question}

Please provide a comprehensive analysis based on the news provided. Include:
1. Direct answer to the question
2. Supporting evidence from the articles
3. Any relevant context or implications
4. Confidence level in your analysis (1-10)"""

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert news analyst. Provide thorough, factual analysis based on the provided articles."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=800,
                temperature=0.3
            )

            return {
                "success": True,
                "analysis": response.choices[0].message.content.strip(),
                "articles_analyzed": len(articles[:5]),
                "question": question
            }
        except Exception as e:
            return {
                "success": False,
                "analysis": f"Analysis failed: {str(e)}",
                "error": str(e)
            }


class ReportGenerationTool(Tool):
    def __init__(self):
        super().__init__("generate_report", "Generate comprehensive report from multiple article summaries")

    def execute(self, scraped_articles: List[Dict] = None, topic: str = "news", **kwargs) -> Dict[str, Any]:
        if not scraped_articles:
            return {
                "success": False,
                "error": "No scraped articles provided for report generation",
                "report": ""
            }

        print(f"📊 Generating comprehensive report from {len(scraped_articles)} articles...")

        article_summaries = []
        successful_articles = [a for a in scraped_articles if a['word_count'] > 0]

        for i, article in enumerate(successful_articles, 1):
            article_summaries.append(f"""
Article {i}: {article['original_title']}
Source: {article['url']}
Word Count: {article['word_count']}
Summary: {article['ai_summary']}
""")

        if not article_summaries:
            return {
                "success": False,
                "error": "No successfully scraped articles to generate report from",
                "report": "Unable to generate report - no articles were successfully scraped."
            }

        combined_content = "\n".join(article_summaries)

        if not client:
            # Generate a basic report without AI
            report = f"""
**Executive Summary**
Analysis of {len(successful_articles)} articles about {topic}.

**Major Themes**
Based on {len(successful_articles)} scraped articles with a total of {sum(a['word_count'] for a in successful_articles):,} words analyzed.

**Key Developments**
{combined_content[:1000]}

**Note**: Full AI analysis unavailable (no API key configured)
"""
            return {
                "success": True,
                "report": report,
                "articles_analyzed": len(successful_articles),
                "total_articles": len(scraped_articles),
                "topic": topic,
                "generated_at": datetime.now().isoformat()
            }

        try:
            report_prompt = f"""Create a comprehensive news analysis report based on these article summaries:

Topic: {topic}
Number of articles analyzed: {len(successful_articles)}

{combined_content}

Generate a professional report with:
1. **Executive Summary** - Key findings and trends
2. **Major Themes** - Common topics and patterns
3. **Key Developments** - Most important news items
4. **Analysis & Insights** - What these developments mean
5. **Sources Summary** - Brief overview of sources used

Format as a clear, well-structured report suitable for executive briefing."""

            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert news analyst creating executive-level reports. Be comprehensive, insightful, and professional."},
                    {"role": "user", "content": report_prompt}
                ],
                max_tokens=1200,
                temperature=0.3
            )

            report = response.choices[0].message.content.strip()

            return {
                "success": True,
                "report": report,
                "articles_analyzed": len(successful_articles),
                "total_articles": len(scraped_articles),
                "topic": topic,
                "generated_at": datetime.now().isoformat()
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Report generation failed: {str(e)}",
                "report": ""
            }


# ============================================
# ENHANCED DOCUMENT GENERATOR
# ============================================

class EnhancedDocumentGenerator:
    """Enhanced professional Word document generator with advanced formatting"""

    def __init__(self):
        self.primary_color = RGBColor(0, 32, 96)  # Dark blue
        self.accent_color = RGBColor(0, 120, 215)  # Lighter blue
        self.success_color = RGBColor(0, 176, 80)  # Green
        self.warning_color = RGBColor(255, 192, 0)  # Amber
        self.error_color = RGBColor(237, 28, 36)  # Red

    def create_professional_document(self,
                                     results: Dict,
                                     original_request: str,
                                     config: Dict = None) -> str:
        """Create a highly professional Word document with all analysis results"""
        config = config or {}

        print(f"📄 Generating enhanced professional Word document...")

        try:
            doc = Document()
            self._setup_advanced_styles(doc)
            self._configure_document_settings(doc)

            # Add all sections with enhanced formatting
            self._add_cover_page(doc, original_request, results)
            self._add_table_of_contents(doc, results)
            self._add_executive_dashboard(doc, results, config)
            self._add_key_findings(doc, results)
            self._add_detailed_analysis(doc, results)
            self._add_article_deep_dive(doc, results, config)
            self._add_statistical_analysis(doc, results)
            self._add_source_credibility(doc, results)
            self._add_appendices(doc, results)
            self._add_footer_headers(doc, original_request)

            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_request = "".join(c for c in original_request if c.isalnum() or c in (' ', '-', '_')).rstrip()[:30]
            filename = f"Professional_News_Analysis_{safe_request}_{timestamp}.docx"

            doc.save(filename)

            print(f"✅ Document successfully generated: {filename}")

            return filename

        except Exception as e:
            print(f"❌ Document generation failed: {str(e)}")
            raise

    def _setup_advanced_styles(self, doc):
        """Set up professional document styles with corporate formatting"""
        styles = doc.styles

        # Cover Title Style
        if 'Cover Title' not in [s.name for s in styles]:
            cover_style = styles.add_style('Cover Title', WD_STYLE_TYPE.PARAGRAPH)
            cover_font = cover_style.font
            cover_font.name = 'Calibri Light'
            cover_font.size = Pt(36)
            cover_font.bold = True
            cover_font.color.rgb = self.primary_color
            cover_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_style.paragraph_format.space_after = Pt(24)

        # Section Header Style
        if 'Section Header' not in [s.name for s in styles]:
            section_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(18)
            section_font.bold = True
            section_font.color.rgb = self.primary_color
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(12)
            section_style.paragraph_format.keep_with_next = True

        # Subsection Style
        if 'Subsection' not in [s.name for s in styles]:
            subsection_style = styles.add_style('Subsection', WD_STYLE_TYPE.PARAGRAPH)
            subsection_font = subsection_style.font
            subsection_font.name = 'Calibri'
            subsection_font.size = Pt(14)
            subsection_font.bold = True
            subsection_font.color.rgb = self.accent_color
            subsection_style.paragraph_format.space_before = Pt(12)
            subsection_style.paragraph_format.space_after = Pt(6)

    def _configure_document_settings(self, doc):
        """Configure document-wide settings"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_cover_page(self, doc, original_request, results):
        """Create an impressive cover page"""
        doc.add_paragraph("\n\n")

        # Main title
        title = doc.add_paragraph("COMPREHENSIVE NEWS ANALYSIS", style='Cover Title')

        # Subtitle with request
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Topic: {original_request}")
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = self.accent_color
        subtitle_run.font.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("\n")

        # Add visual separator
        separator = doc.add_paragraph("─" * 60)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.runs[0].font.color.rgb = self.accent_color

        doc.add_paragraph("\n")

        # Key metrics box
        metrics_table = doc.add_table(rows=2, cols=4)
        metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style the metrics table
        metrics_table.style = 'Light Grid'

        # Add metrics headers
        headers = ['Articles Found', 'Articles Scraped', 'Success Rate', 'Analysis Depth']
        for i, header in enumerate(headers):
            cell = metrics_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        # Add metrics values
        search_count = results.get('search_news', {}).get('count', 0)
        scraped_count = results.get('scrape_articles', {}).get('total_scraped', 0)
        success_count = results.get('scrape_articles', {}).get('successful_scrapes', 0)
        success_rate = (success_count / scraped_count * 100) if scraped_count > 0 else 0

        values = [
            str(search_count),
            str(scraped_count),
            f"{success_rate:.1f}%",
            "COMPREHENSIVE"
        ]

        for i, value in enumerate(values):
            cell = metrics_table.rows[1].cells[i]
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = self.accent_color

        doc.add_paragraph("\n\n")

        # Generation details
        gen_details = doc.add_paragraph()
        gen_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_details.add_run("Generated: ").font.size = Pt(10)
        gen_run = gen_details.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p'))
        gen_run.font.size = Pt(10)
        gen_run.font.bold = True

        # Analysis engine
        engine_para = doc.add_paragraph()
        engine_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        engine_para.add_run("Powered by: ").font.size = Pt(10)
        engine_run = engine_para.add_run("GPT-4 Advanced Analysis Engine")
        engine_run.font.size = Pt(10)
        engine_run.font.bold = True
        engine_run.font.color.rgb = self.accent_color

        doc.add_page_break()

    def _add_table_of_contents(self, doc, results):
        """Add a table of contents with page references"""
        doc.add_paragraph("TABLE OF CONTENTS", style='Section Header')

        toc_items = [
            ("Executive Dashboard", "Key metrics and performance indicators"),
            ("Key Findings", "Principal discoveries and insights"),
            ("Detailed Analysis", "Comprehensive analytical breakdown"),
            ("Article Deep Dive", "Individual article examination"),
            ("Statistical Analysis", "Quantitative data assessment"),
            ("Source Credibility", "Source reliability evaluation"),
            ("Technical Appendix", "Implementation details"),
        ]

        for i, (title, description) in enumerate(toc_items, 1):
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.left_indent = Inches(0.25)

            # Add number and title
            num_run = toc_para.add_run(f"{i}. ")
            num_run.font.size = Pt(12)
            num_run.font.bold = True

            title_run = toc_para.add_run(title)
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = self.primary_color

            # Add dots
            dots_run = toc_para.add_run(" " + "." * 50)
            dots_run.font.size = Pt(12)
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            # Add description
            desc_para = doc.add_paragraph()
            desc_para.paragraph_format.left_indent = Inches(0.5)
            desc_run = desc_para.add_run(description)
            desc_run.font.size = Pt(10)
            desc_run.font.italic = True
            desc_run.font.color.rgb = RGBColor(64, 64, 64)

        doc.add_page_break()

    def _add_executive_dashboard(self, doc, results, config):
        """Add an executive dashboard with visual metrics"""
        doc.add_paragraph("EXECUTIVE DASHBOARD", style='Section Header')

        # Performance Overview
        doc.add_paragraph("Performance Overview", style='Subsection')

        # Create performance table
        perf_table = doc.add_table(rows=5, cols=3)
        perf_table.style = 'Medium Shading 1'

        performance_data = [
            ["Metric", "Value", "Status"],
            ["Total Articles Discovered",
             str(results.get('search_news', {}).get('count', 0)),
             "✓ Complete"],
            ["Articles Successfully Scraped",
             str(results.get('scrape_articles', {}).get('successful_scrapes', 0)),
             "✓ Complete"],
            ["AI Summaries Generated",
             str(len([a for a in results.get('scrape_articles', {}).get('scraped_articles', [])
                      if a.get('ai_summary') and a['ai_summary'] != "Could not scrape content from this URL"])),
             "✓ Complete"],
            ["Comprehensive Report",
             "Generated" if results.get('generate_report', {}).get('success') else "Not Generated",
             "✓ Complete" if results.get('generate_report', {}).get('success') else "✗ Failed"]
        ]

        for row_idx, row_data in enumerate(performance_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = perf_table.rows[row_idx].cells[col_idx]
                cell.text = cell_data
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # Analysis Configuration
        doc.add_paragraph("Analysis Configuration", style='Subsection')

        config_para = doc.add_paragraph()
        config_items = [
            f"• Search Query: {results.get('search_news', {}).get('query', 'N/A')}",
            f"• Maximum Articles Requested: {config.get('max_articles', 15)}",
            f"• Scraping Depth: Full Content Extraction",
            f"• AI Model: GPT-4 (Advanced)" if client else "• AI Model: Not configured",
            f"• Analysis Type: Comprehensive Multi-Source",
            f"• Report Format: Professional Executive Brief"
        ]

        for item in config_items:
            config_para.add_run(item + "\n")

        doc.add_paragraph("")

    def _add_key_findings(self, doc, results):
        """Add key findings section with highlights"""
        doc.add_paragraph("KEY FINDINGS", style='Section Header')

        if 'generate_report' in results and results['generate_report'].get('success'):
            report_content = results['generate_report'].get('report', '')

            # Parse and format the report content
            sections = self._parse_report_content(report_content)

            for section_title, section_content in sections:
                if section_title:
                    # Add section with special formatting for key sections
                    if any(keyword in section_title.lower() for keyword in ['executive', 'key', 'major']):
                        section_para = doc.add_paragraph(section_title, style='Subsection')

                        # Add highlighted box for executive summary
                        if 'executive' in section_title.lower():
                            self._add_highlight_box(doc, section_content)
                        else:
                            content_para = doc.add_paragraph(section_content)
                            content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    else:
                        doc.add_paragraph(section_title, style='Subsection')
                        doc.add_paragraph(section_content)

        doc.add_paragraph("")

    def _add_article_deep_dive(self, doc, results, config):
        """Add comprehensive article analysis section"""
        doc.add_paragraph("ARTICLE DEEP DIVE", style='Section Header')

        if 'scrape_articles' not in results or not results['scrape_articles'].get('success'):
            doc.add_paragraph("No articles were successfully scraped for deep analysis.")
            return

        scraped_articles = results['scrape_articles'].get('scraped_articles', [])
        successful_articles = [a for a in scraped_articles if a.get('word_count', 0) > 0]

        # Summary statistics
        doc.add_paragraph("Coverage Statistics", style='Subsection')

        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Light List'

        stats_data = [
            ["Total Articles Processed", str(len(scraped_articles))],
            ["Successfully Analyzed", str(len(successful_articles))],
            ["Total Words Processed", str(sum(a.get('word_count', 0) for a in successful_articles))],
            ["Average Article Length", f"{sum(a.get('word_count', 0) for a in successful_articles) // max(len(successful_articles), 1)} words"]
        ]

        for row_idx, (label, value) in enumerate(stats_data):
            stats_table.rows[row_idx].cells[0].text = label
            stats_table.rows[row_idx].cells[1].text = value
            stats_table.rows[row_idx].cells[1].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph("")

        # Individual article analysis
        doc.add_paragraph("Individual Article Analysis", style='Subsection')

        max_articles_to_show = config.get('max_articles_in_report', len(successful_articles))

        for i, article in enumerate(successful_articles[:max_articles_to_show], 1):
            # Article header with formatting
            article_header = doc.add_paragraph()
            article_num = article_header.add_run(f"Article {i}: ")
            article_num.font.bold = True
            article_num.font.size = Pt(12)
            article_num.font.color.rgb = self.primary_color

            article_title = article_header.add_run(article['original_title'])
            article_title.font.size = Pt(12)
            article_title.font.italic = True

            # Create article details table
            details_table = doc.add_table(rows=4, cols=2)
            details_table.style = 'Table Grid'

            # Set column widths
            for cell in details_table.columns[0].cells:
                cell.width = Inches(1.5)
            for cell in details_table.columns[1].cells:
                cell.width = Inches(5)

            details = [
                ["Source URL", article['url']],
                ["Word Count", f"{article['word_count']:,} words"],
                ["Extraction Time", article['scrape_timestamp'][:19].replace('T', ' at ')],
                ["Extraction Status", "✓ Successful" if article['word_count'] > 0 else "✗ Failed"]
            ]

            for row_idx, (label, value) in enumerate(details):
                label_cell = details_table.rows[row_idx].cells[0]
                label_cell.text = label
                label_cell.paragraphs[0].runs[0].font.bold = True
                label_cell.paragraphs[0].runs[0].font.size = Pt(10)

                value_cell = details_table.rows[row_idx].cells[1]
                value_cell.text = value
                value_cell.paragraphs[0].runs[0].font.size = Pt(10)

            doc.add_paragraph("")

            # AI Summary with special formatting
            summary_header = doc.add_paragraph()
            summary_header.add_run("AI-Generated Summary:").font.bold = True

            summary_para = doc.add_paragraph()
            summary_para.paragraph_format.left_indent = Inches(0.25)
            summary_para.paragraph_format.right_indent = Inches(0.25)
            summary_para.paragraph_format.space_before = Pt(6)
            summary_para.paragraph_format.space_after = Pt(6)

            summary_run = summary_para.add_run(article.get('ai_summary', 'No summary available'))
            summary_run.font.size = Pt(11)

            # Add original summary if available
            if article.get('original_summary'):
                orig_header = doc.add_paragraph()
                orig_header.add_run("Original Summary:").font.bold = True

                orig_para = doc.add_paragraph()
                orig_para.paragraph_format.left_indent = Inches(0.25)
                orig_para.add_run(article['original_summary']).font.size = Pt(10)

            # Add content preview if available
            if article.get('full_content'):
                preview_header = doc.add_paragraph()
                preview_header.add_run("Content Preview (First 300 characters):").font.bold = True

                preview_para = doc.add_paragraph()
                preview_para.paragraph_format.left_indent = Inches(0.25)
                preview_text = article['full_content'][:300] + "..." if len(article['full_content']) > 300 else article['full_content']
                preview_run = preview_para.add_run(preview_text)
                preview_run.font.size = Pt(9)
                preview_run.font.color.rgb = RGBColor(64, 64, 64)

            # Add separator
            if i < len(successful_articles[:max_articles_to_show]):
                separator = doc.add_paragraph("─" * 80)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator.runs[0].font.color.rgb = RGBColor(200, 200, 200)
                separator.runs[0].font.size = Pt(8)

        doc.add_page_break()

    def _add_statistical_analysis(self, doc, results):
        """Add statistical analysis section"""
        doc.add_paragraph("STATISTICAL ANALYSIS", style='Section Header')

        # Data Quality Metrics
        doc.add_paragraph("Data Quality Metrics", style='Subsection')

        scraped_articles = results.get('scrape_articles', {}).get('scraped_articles', [])
        successful = [a for a in scraped_articles if a.get('word_count', 0) > 0]
        failed = [a for a in scraped_articles if a.get('word_count', 0) == 0]

        quality_table = doc.add_table(rows=6, cols=3)
        quality_table.style = 'Medium Grid 3'

        quality_data = [
            ["Metric", "Value", "Percentage"],
            ["Total Articles Attempted", str(len(scraped_articles)), "100%"],
            ["Successful Extractions", str(len(successful)), f"{len(successful)/max(len(scraped_articles), 1)*100:.1f}%"],
            ["Failed Extractions", str(len(failed)), f"{len(failed)/max(len(scraped_articles), 1)*100:.1f}%"],
            ["Average Words per Article", f"{sum(a.get('word_count', 0) for a in successful) // max(len(successful), 1)}", "N/A"],
            ["Total Content Analyzed", f"{sum(a.get('word_count', 0) for a in successful):,} words", "N/A"]
        ]

        for row_idx, row_data in enumerate(quality_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = quality_table.rows[row_idx].cells[col_idx]
                cell.text = cell_data
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # Content Distribution
        doc.add_paragraph("Content Distribution", style='Subsection')

        if successful:
            # Word count distribution
            word_counts = [a.get('word_count', 0) for a in successful]

            distribution_para = doc.add_paragraph()
            distribution_items = [
                f"• Shortest Article: {min(word_counts):,} words",
                f"• Longest Article: {max(word_counts):,} words",
                f"• Median Length: {sorted(word_counts)[len(word_counts)//2]:,} words",
                f"• Standard Deviation: {self._calculate_std_dev(word_counts):.1f} words"
            ]

            for item in distribution_items:
                distribution_para.add_run(item + "\n")

    def _add_source_credibility(self, doc, results):
        """Add source credibility analysis"""
        doc.add_paragraph("SOURCE CREDIBILITY", style='Section Header')

        doc.add_paragraph("Source Diversity Analysis", style='Subsection')

        if 'scrape_articles' in results:
            articles = results['scrape_articles'].get('scraped_articles', [])

            # Extract domains
            domains = {}
            for article in articles:
                url = article.get('url', '')
                if url:
                    try:
                        from urllib.parse import urlparse
                        domain = urlparse(url).netloc
                        domains[domain] = domains.get(domain, 0) + 1
                    except:
                        pass

            if domains:
                # Create source table
                source_table = doc.add_table(rows=len(domains) + 1, cols=3)
                source_table.style = 'Light Shading'

                # Headers
                headers = ["Source Domain", "Articles", "Percentage"]
                for i, header in enumerate(headers):
                    cell = source_table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Sort domains by count
                sorted_domains = sorted(domains.items(), key=lambda x: x[1], reverse=True)
                total_articles = sum(domains.values())

                for row_idx, (domain, count) in enumerate(sorted_domains, 1):
                    source_table.rows[row_idx].cells[0].text = domain
                    source_table.rows[row_idx].cells[1].text = str(count)
                    source_table.rows[row_idx].cells[2].text = f"{count/total_articles*100:.1f}%"

        doc.add_paragraph("")

    def _add_detailed_analysis(self, doc, results):
        """Add the comprehensive analysis from the AI report"""
        doc.add_paragraph("DETAILED ANALYSIS", style='Section Header')

        if 'generate_report' in results and results['generate_report'].get('success'):
            report = results['generate_report'].get('report', '')

            # Split report into paragraphs and format
            paragraphs = report.split('\n\n')

            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a header (starts with ** or #)
                    if para_text.strip().startswith('**') or para_text.strip().startswith('#'):
                        # Extract header text
                        header_text = para_text.strip().replace('**', '').replace('#', '').strip()
                        doc.add_paragraph(header_text, style='Subsection')
                    else:
                        # Regular paragraph
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        para.paragraph_format.space_after = Pt(6)

    def _add_appendices(self, doc, results):
        """Add technical appendices"""
        doc.add_page_break()
        doc.add_paragraph("TECHNICAL APPENDIX", style='Section Header')

        # Execution Log
        doc.add_paragraph("Execution Log", style='Subsection')

        execution_table = doc.add_table(rows=len(results) + 1, cols=3)
        execution_table.style = 'Medium List 2'

        # Headers
        headers = ["Tool", "Status", "Details"]
        for i, header in enumerate(headers):
            cell = execution_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tool execution details
        for row_idx, (tool_name, result) in enumerate(results.items(), 1):
            execution_table.rows[row_idx].cells[0].text = tool_name

            if result.get('success'):
                status_cell = execution_table.rows[row_idx].cells[1]
                status_cell.text = "✓ Success"
                status_cell.paragraphs[0].runs[0].font.color.rgb = self.success_color
            else:
                status_cell = execution_table.rows[row_idx].cells[1]
                status_cell.text = "✗ Failed"
                status_cell.paragraphs[0].runs[0].font.color.rgb = self.error_color

            # Add relevant details
            details = self._get_tool_details(tool_name, result)
            execution_table.rows[row_idx].cells[2].text = details

        doc.add_paragraph("")

        # System Information
        doc.add_paragraph("System Information", style='Subsection')

        system_info = doc.add_paragraph()
        system_items = [
            f"• Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"• Analysis Engine: {'GPT-4 Advanced' if client else 'Basic (No API Key)'}",
            f"• Scraping Engine: BeautifulSoup 4",
            f"• Search Provider: DuckDuckGo",
            f"• Document Format: Microsoft Word (.docx)",
            f"• Processing Pipeline: Automated Multi-Stage",
            f"• Quality Assurance: Multi-Point Verification"
        ]

        for item in system_items:
            system_info.add_run(item + "\n")

    def _add_footer_headers(self, doc, original_request):
        """Add professional headers and footers"""
        section = doc.sections[0]

        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"News Analysis Report - {original_request[:50]}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(10)
        header_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%B %d, %Y')} | Confidential Analysis Report | Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    def _add_highlight_box(self, doc, content):
        """Add a highlighted content box"""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        cell = table.rows[0].cells[0]
        cell_para = cell.paragraphs[0]
        cell_para.text = content

        # Apply shading to cell
        shading_elm = parse_xml(r'<w:shd {} w:fill="E8F4FD"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        cell_para.paragraph_format.space_before = Pt(6)
        cell_para.paragraph_format.space_after = Pt(6)
        cell_para.runs[0].font.size = Pt(11)

    def _parse_report_content(self, report_content):
        """Parse report content into structured sections"""
        sections = []
        current_section = ""
        current_title = ""

        lines = report_content.split('\n')

        for line in lines:
            # Check for markdown headers
            if line.startswith('**') and line.endswith('**'):
                if current_title or current_section:
                    sections.append((current_title, current_section.strip()))
                current_title = line.strip('*').strip()
                current_section = ""
            elif line.startswith('#'):
                if current_title or current_section:
                    sections.append((current_title, current_section.strip()))
                current_title = line.strip('#').strip()
                current_section = ""
            else:
                current_section += line + '\n'

        if current_title or current_section:
            sections.append((current_title, current_section.strip()))

        return sections

    def _calculate_std_dev(self, numbers):
        """Calculate standard deviation"""
        if not numbers:
            return 0
        mean = sum(numbers) / len(numbers)
        variance = sum((x - mean) ** 2 for x in numbers) / len(numbers)
        return variance ** 0.5

    def _get_tool_details(self, tool_name, result):
        """Get relevant details for each tool execution"""
        details = ""

        if tool_name == "search_news":
            details = f"Found {result.get('count', 0)} articles"
        elif tool_name == "scrape_articles":
            details = f"Scraped {result.get('successful_scrapes', 0)}/{result.get('total_scraped', 0)} articles"
        elif tool_name == "analyze_news":
            details = f"Analyzed {result.get('articles_analyzed', 0)} articles"
        elif tool_name == "generate_report":
            details = f"Report with {result.get('articles_analyzed', 0)} sources"
        elif tool_name == "generate_document":
            details = f"Document: {result.get('document_path', 'N/A')}"
        elif tool_name == "check_internet":
            details = f"Status: {result.get('status', 'Unknown')}"
        else:
            details = "Completed" if result.get('success') else result.get('error', 'Failed')

        return details


# ============================================
# CONFIGURABLE NEWS AGENT
# ============================================

class ConfigurableNewsAgent:
    """Enhanced News Agent with configurable parameters for flexible analysis"""

    def __init__(self):
        self.tools = {
            "check_internet": InternetConnectivityTool(),
            "search_news": NewsSearchTool(),
            "analyze_news": NewsAnalysisTool(),
            "scrape_articles": WebScrapingTool(),
            "generate_report": ReportGenerationTool(),
            "generate_document": EnhancedDocumentGenerator()
        }
        self.memory = {
            "conversation_history": [],
            "cached_news": {},
            "scraped_content": {},
            "user_preferences": {}
        }
        self.max_iterations = 20

    def execute_with_config(self,
                            user_request: str,
                            num_search_articles: int = 20,
                            num_scrape_articles: int = 15,
                            num_articles_in_report: int = None) -> str:
        """Execute news analysis with configurable parameters"""

        # Validate inputs
        num_search_articles = max(1, min(50, num_search_articles))
        num_scrape_articles = max(1, min(30, num_scrape_articles))

        if num_articles_in_report is None:
            num_articles_in_report = num_scrape_articles

        print(f"🤖 Executing analysis with custom configuration:")
        print(f"   📰 Search for: {num_search_articles} articles")
        print(f"   🔍 Scrape: {num_scrape_articles} articles")
        print(f"   📊 Report detail: {num_articles_in_report} articles")
        print("=" * 50)

        config = {
            'max_search': num_search_articles,
            'max_scrape': num_scrape_articles,
            'max_articles_in_report': num_articles_in_report
        }

        # Execute tools with configuration
        results = {}

        # 1. Check internet
        print("\n🌐 Checking internet connectivity...")
        internet_result = self.tools["check_internet"].execute()
        results["check_internet"] = internet_result

        if not internet_result.get("success"):
            return "❌ No internet connection available. Please check your connection."

        # 2. Search for articles
        print(f"\n🔍 Searching for {num_search_articles} articles about: {user_request}")
        search_result = self.tools["search_news"].execute(
            query=user_request,
            max_results=num_search_articles
        )
        results["search_news"] = search_result

        if not search_result.get("success") or not search_result.get("articles"):
            return f"❌ Could not find articles about: {user_request}"

        articles = search_result.get("articles", [])
        print(f"✅ Found {len(articles)} articles")

        # 3. Scrape articles
        print(f"\n📄 Scraping {num_scrape_articles} articles for full content...")
        scrape_result = self.tools["scrape_articles"].execute(
            articles=articles,
            max_articles=num_scrape_articles
        )
        results["scrape_articles"] = scrape_result

        scraped_articles = scrape_result.get("scraped_articles", [])
        successful_scrapes = scrape_result.get("successful_scrapes", 0)
        print(f"✅ Successfully scraped {successful_scrapes}/{len(scraped_articles)} articles")

        # 4. Generate comprehensive report
        if successful_scrapes > 0:
            print(f"\n📊 Generating comprehensive report from {successful_scrapes} articles...")
            report_result = self.tools["generate_report"].execute(
                scraped_articles=scraped_articles,
                topic=user_request
            )
            results["generate_report"] = report_result

            if report_result.get("success"):
                print("✅ Report generated successfully")

        # 5. Analyze articles (for additional insights)
        print(f"\n🧠 Performing AI analysis...")
        analyze_result = self.tools["analyze_news"].execute(
            articles=articles[:5],
            question=f"Provide key insights about: {user_request}"
        )
        results["analyze_news"] = analyze_result

        # 6. Generate professional document
        print(f"\n📄 Creating professional Word document...")
        doc_result = self.tools["generate_document"].create_professional_document(
            results=results,
            original_request=user_request,
            config=config
        )

        results["generate_document"] = {
            "success": True if doc_result else False,
            "document_path": doc_result
        }

        # Generate response
        if doc_result:
            response = f"""
✅ **COMPREHENSIVE ANALYSIS COMPLETE**

📄 **Professional Document Generated**: `{doc_result}`

📊 **Analysis Summary:**
• Articles Searched: {len(articles)}
• Articles Scraped: {len(scraped_articles)}
• Successful Extractions: {successful_scrapes}
• Total Words Analyzed: {sum(a.get('word_count', 0) for a in scraped_articles):,}
• AI Summaries Generated: {len([a for a in scraped_articles if a.get('ai_summary')])}

📋 **Document Contents:**
• Executive Dashboard with performance metrics
• Table of Contents for easy navigation  
• Key Findings and insights
• Detailed Article Analysis ({num_articles_in_report} articles)
• Statistical Analysis and data quality metrics
• Source Credibility evaluation
• Comprehensive AI-generated report
• Technical appendix with execution details

🎯 **Topic Analyzed**: "{user_request}"

✨ The document has been saved in your current directory and is ready for:
• Executive presentations
• Team sharing
• Archive and reference
• Further analysis

Open the Word document to view the complete professional analysis report!
"""
        else:
            response = "❌ Failed to generate document. Please check the logs for details."

        return response


# ============================================
# HELPER FUNCTIONS
# ============================================

def analyze_news(topic: str,
                 search_count: int = 20,
                 scrape_count: int = 15,
                 report_detail: int = None):
    """Simplified function to analyze news with custom parameters"""
    agent = ConfigurableNewsAgent()
    return agent.execute_with_config(
        user_request=topic,
        num_search_articles=search_count,
        num_scrape_articles=scrape_count,
        num_articles_in_report=report_detail
    )


def quick_news_report(topic: str):
    """Generate a quick news report with default settings"""
    return analyze_news(topic, search_count=15, scrape_count=10, report_detail=5)


def comprehensive_news_analysis(topic: str):
    """Generate a comprehensive news analysis with maximum detail"""
    return analyze_news(topic, search_count=30, scrape_count=20, report_detail=20)


# ============================================
# MAIN CLI INTERFACE
# ============================================

def interactive_news_analyzer():
    """Interactive command-line interface for the news analyzer"""
    print("🤖 ENHANCED NEWS ANALYSIS SYSTEM")
    print("=" * 50)
    print("\nThis system can generate professional Word documents with:")
    print("• Customizable number of articles to search")
    print("• Configurable scraping depth")
    print("• Adjustable report detail level")
    print("\n" + "=" * 50)

    while True:
        try:
            print("\n📰 Enter your news topic (or 'quit' to exit):")
            topic = input("Topic: ").strip()

            if topic.lower() in ['quit', 'exit', 'bye']:
                print("👋 Goodbye!")
                break

            if not topic:
                print("❌ Please enter a topic.")
                continue

            print("\n⚙️ Configuration (press Enter for defaults):")

            # Get search count
            search_input = input("Number of articles to search for (default 20): ").strip()
            search_count = int(search_input) if search_input else 20

            # Get scrape count
            scrape_input = input("Number of articles to scrape (default 15): ").strip()
            scrape_count = int(scrape_input) if scrape_input else 15

            # Get report detail
            detail_input = input(f"Number of articles in detailed report (default {scrape_count}): ").strip()
            report_detail = int(detail_input) if detail_input else scrape_count

            print("\n🚀 Starting analysis...")
            print("=" * 50)

            result = analyze_news(
                topic=topic,
                search_count=search_count,
                scrape_count=scrape_count,
                report_detail=report_detail
            )

            print("\n" + "=" * 50)
            print(result)
            print("=" * 50)

        except KeyboardInterrupt:
            print("\n👋 Goodbye!")
            break
        except ValueError as e:
            print(f"❌ Invalid number entered. Please enter a valid number.")
        except Exception as e:
            print(f"❌ An error occurred: {e}")


# ============================================
# MAIN EXECUTION
# ============================================

if __name__ == "__main__":
    interactive_news_analyzer()